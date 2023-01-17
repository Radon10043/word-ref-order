'''
Author: Radon
Date: 2023-01-16 15:48:19
LastEditors: Radon
LastEditTime: 2023-01-17 09:57:44
Description: Hi, say something
'''
import docx
import re
import argparse
import sys, os

REF_REGEXP = "\[[0-9]+\]"  # 用于判断一段文本是否为参考文献引用的正则表达式


def check_valid(main_file: str, ref_file: str) -> bool:
    """检查输入的合法性

    Parameters
    ----------
    main_file : str
        正文文件路径
    ref_file : str
        参考文献文件路径

    Returns
    -------
    bool
        True : 输入合法
        False : 输入不合法

    Notes
    -----
    _description_
    """
    if not os.path.exists(main_file):
        print("正文文件" + main_file + "不存在!")
        return False

    if not os.path.exists(ref_file):
        print("参考文献文件" + ref_file + "不存在!")
        return False

    if os.path.splitext(main_file)[-1] != ".docx":
        print("正文文件" + main_file + "不是.docx文件!")
        return False

    if os.path.splitext(ref_file)[-1] != ".docx":
        print("参考文献文件" + ref_file + "不是.docx文件!")
        return False

    return True


def main(main_file: str, ref_file: str):
    """整理docx中交叉引用的编号

    Parameters
    ----------
    main_file : str
        正文文件路径
    ref_file : str
        参考文献文件路径

    Notes
    -----
    _description_
    """
    # 读取参考文献文件, 将其中的参考文献加入ref_lst, 文件中的参考文献需要每个都编号
    ref_lst = list()  # 参考文献列表
    ref_set = set()  # 参考文献集合, 用于去重
    rep_set = set()  # 重复的参考文献集合
    f = docx.Document(ref_file)
    for paragraph in f.paragraphs:
        ref_name = paragraph.text
        if paragraph.style.name == "List Paragraph":  # 当前段落被编号的话才会被认为是一个参考文献
            if ref_name in ref_set:
                rep_set.add(ref_name)
            else:
                ref_set.add(ref_name)
                ref_lst.append(ref_name)

    # 如果有重复参考文献, 提示并退出
    if len(rep_set):
        print("哎呀! 看起来这些参考文献在列表中出现了多次:\n")
        for ref in rep_set:
            print(ref)
        print("\n请去重更新后再重试!")
        return

    # 遍历正文, 记录有交叉引用过文献的地方
    cr_lst = list()  # cross references list, 根据出现的顺序记录每个交叉引用的序号
    f = docx.Document(main_file)
    for paragraph in f.paragraphs:
        for run in paragraph.runs:
            # FIXME: python-docx中似乎没有能识别交叉引用的功能, 这里用的是正则表达式
            # 暂时不清楚是否存在BUG
            is_cf = re.fullmatch(REF_REGEXP, run.text)
            if is_cf:
                cr_num = int(run.text.lstrip("[").rstrip("]"))
                cr_lst.append(cr_num - 1)  # 将交叉引用的序号加入cr_lst, 这里减1是为了能和ref_lst的下标对上

    # 调整参考文献顺序
    cr_set = set()  # cross references set, 记录存储过的交叉引用集合
    nref_lst = list()  # new references list, 调整好顺序后的参考文献列表
    for num in cr_lst:
        # 根据出现顺序将ref加入nref_lst, 如果已经出现过了就不再添加
        if num in cr_set: continue
        cr_set.add(num)
        nref_lst.append(ref_lst[num])

    # 输出
    nfn = "_new".join(os.path.splitext(ref_file))  # New file name
    nf = docx.Document()
    for ncr in nref_lst:
        nf.add_paragraph(ncr)
    nf.save(nfn)

    # 完成!
    print("完成! 请查看" + nfn + "!")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("-m", "--main", help="存储正文内容与参考文献内容的文件", required=True)
    parser.add_argument("-r", "--ref", help="仅存储参考文献内容文件, 且参考文献已编号", required=True)
    args = parser.parse_args()
    main_file, ref_file = args.main, args.ref

    if not check_valid(main_file, ref_file):
        sys.exit(0)

    main(main_file, ref_file)